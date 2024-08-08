from docx import Document
from docx.shared import Inches
import os

def save_file(doc, filename):
    # Salva o arquivo em word e pdf
    if os.path.isdir(f'{os.getcwd()}/word'):
        doc.save(f'{os.getcwd()}/word/{filename}.docx')
    else:
        os.mkdir('word')
        doc.save(f'{os.getcwd()}/word/{filename}.docx')


def date_formart(date:str):
    date = date.split("-")
    return f'{date[0]}/{date[1]}/{date[2]}'


def create_word(coin, date) -> str:
    site = 'https://wise.com/br/currency-converter/usd-to-brl-rate?amount=1'
    name = 'João Matheus Albergaria Rodrigues Santana'
    
    # Cria o Documento
    doc = Document()
    
    # Style doc
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    
    # Adiciona um titulo
    doc.add_heading(f"Cotação Atual do Dólar - R${coin} ({date_formart(date)})", 0)
    
    # Adiciona paragrafos
    doc.add_paragraph(f'O dolár está no valor de R${coin} na data ({date_formart(date)})')
    doc.add_paragraph(f'Valor cotado no site {site}')
    doc.add_paragraph('Print da cotação atual')
    
    # Adiciona Imagem
    doc.add_picture(f'screenshot/{date}.png',width=Inches(6.25))
    
    # Adiciona paragrafo
    doc.add_paragraph(f'Cotação feita por - {name}')
    
    # Salva o Arquivo
    save_file(doc, date)
    return date